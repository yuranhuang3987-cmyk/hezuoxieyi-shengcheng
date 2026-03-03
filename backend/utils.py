# -*- coding: utf-8 -*-
"""协议生成工具函数"""

import os
import re
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy


def extract_info(file_path):
    """
    从 .docx 文件中提取软件著作权信息
    使用与原版桌面程序完全一致的逻辑

    Args:
        file_path: .docx 文件路径

    Returns:
        list: 软件信息列表，每个元素是一个 dict
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tables = []

    with zipfile.ZipFile(file_path) as z:
        xml = z.read("word/document.xml")
        root = ET.fromstring(xml)

        for tbl in root.findall(".//w:tbl", ns):
            t = []
            for r in tbl.findall(".//w:tr", ns):
                row = [
                    "".join(x.text or "" for x in c.findall(".//w:t", ns)).strip()
                    for c in r.findall(".//w:tc", ns)
                ]
                if row:
                    t.append(row)
            if t:
                tables.append(t)

    # 提取所有软件信息（支持多软件）
    software_list = []
    owners = []
    software_extracted = False  # 标记是否已提取软件

    for t in tables:
        for i, r in enumerate(t):
            rt = " ".join(r)
            
            # 提取所有软件信息（查找所有"软著名称"行）
            if "软著名称" in rt and i + 1 < len(t) and not software_extracted:
                # 从下一行开始提取所有软件（直到遇到空行或"著作权人信息"）
                for j in range(i + 1, len(t)):
                    d = t[j]
                    # 停止条件：空行或遇到"著作权人"
                    if not d or not d[0] or "著作权人" in d[0]:
                        break
                    
                    # 提取软件信息
                    software_name = d[0] if d else ""
                    if software_name:
                        software_list.append({
                            "name": software_name,
                            "version": d[2] if len(d) > 2 and d[2] else "V1.0",
                            "date": d[3] if len(d) > 3 else "",
                            "owners": []  # 先留空，后面统一填充
                        })
                software_extracted = True  # 标记已提取

            # 提取著作权人信息（所有软件共享）
            if "著作权人信息" in rt:
                for j in range(i + 1, min(i + 3, len(t))):
                    if "公司/单位" in " ".join(t[j]):
                        for k in range(j + 1, len(t)):
                            row = t[k]
                            # 停止条件：遇到联系人、空行
                            if "联系人" in " ".join(row) or not row or not row[0]:
                                break
                            # 跳过标题行（增强过滤）
                            if any(keyword in row[0] for keyword in ["公司", "营业执照", "申请人", "详细地址", "邮政编码"]):
                                continue
                            idn = row[1] if len(row) > 1 else ""
                            owners.append(
                                {
                                    "name": row[0],
                                    "idn": idn.strip(),  # 保存身份证号码
                                    "is_person": bool(re.match(r"^\d{17}[\dXx]$", idn.strip())),
                                }
                            )
                        break

    # 为所有软件填充著作权人信息
    for software in software_list:
        software["owners"] = owners

    # 返回列表格式
    return software_list if software_list else []


def format_date(date_str):
    """
    将日期转换为中文格式（2025年1月1日）
    
    Args:
        date_str: 日期字符串（支持 2025-01-01, 2025/01/01, 2025年1月1日 等格式）
    
    Returns:
        str: 中文格式日期
    """
    try:
        # 已经是中文格式
        if "年" in date_str:
            return date_str
        
        # 处理横杠或斜杠格式
        if "-" in date_str:
            parts = date_str.split("-")
        elif "/" in date_str:
            parts = date_str.split("/")
        else:
            return date_str
        
        if len(parts) >= 3:
            y, m, d = int(parts[0]), int(parts[1]), int(parts[2])
            return f"{y}年{m}月{d}日"
    except:
        pass
    return date_str


def calc_agreement_date(dev_date, custom_date=None):
    """
    根据开发完成日期计算协议签署日期（提前3个月）
    如果传入自定义日期，则直接使用自定义日期

    Args:
        dev_date: 开发完成日期（支持多种格式）
        custom_date: 自定义协议签署日期（可选，格式：2025年1月1日 或 2025-01-01）

    Returns:
        str: 协议签署日期
    """
    # 如果有自定义日期，直接使用
    if custom_date:
        return format_date(custom_date)
    
    # 否则根据开发完成日期计算（提前3个月）
    try:
        # 处理多种日期格式
        if "年" in dev_date:
            # 中文格式：2025年3月10日
            p = dev_date.replace("年", " ").replace("月", " ").replace("日", "").split()
            y, m, d = int(p[0]), int(p[1]) - 3, int(p[2])
        elif "/" in dev_date:
            # 斜杠格式：2025/12/21
            p = dev_date.split("/")
            y, m, d = int(p[0]), int(p[1]) - 3, int(p[2])
        elif "-" in dev_date:
            # 横杠格式：2025-3-10
            p = dev_date.split("-")
            y, m, d = int(p[0]), int(p[1]) - 3, int(p[2])
        else:
            return "2025年1月1日"

        # 处理跨年情况
        if m <= 0:
            m += 12
            y -= 1

        return f"{y}年{m}月{d}日"
    except:
        pass
    return "2025年1月1日"


def check_minor_owners(owners, agreement_date):
    """
    检查是否有未成年著作权人
    
    Args:
        owners: 著作权人列表
        agreement_date: 协议签署日期（格式：2025年1月1日）
    
    Returns:
        list: 未成年著作权人列表 [{"name": "张三", "age": 16}, ...]
    """
    minors = []
    
    try:
        # 解析协议日期
        if "年" in agreement_date:
            p = agreement_date.replace("年", " ").replace("月", " ").replace("日", "").split()
            agr_year, agr_month, agr_day = int(p[0]), int(p[1]), int(p[2])
        else:
            return minors
    except:
        return minors
    
    for owner in owners:
        # 只检查自然人
        if not owner.get("is_person"):
            continue
        
        idn = owner.get("idn", "")
        name = owner.get("name", "")
        
        # 从身份证号码提取出生日期
        if len(idn) >= 14 and idn[:17].isdigit():
            try:
                # 身份证第7-14位是出生日期（YYYYMMDD）
                birth_year = int(idn[6:10])
                birth_month = int(idn[10:12])
                birth_day = int(idn[12:14])
                
                # 计算年龄
                age = agr_year - birth_year
                if (agr_month, agr_day) < (birth_month, birth_day):
                    age -= 1
                
                # 如果未成年，添加到列表
                if age < 18:
                    minors.append({
                        "name": name,
                        "age": age
                    })
            except:
                pass
    
    return minors


def replace_paragraph_text(para, old_text, new_text):
    """
    替换段落中的文本
    
    Args:
        para: 段落对象
        old_text: 要替换的文本
        new_text: 新文本
    """
    if old_text not in para.text:
        return False
    
    runs = para.runs
    if not runs:
        return False
    
    # 获取完整段落文本并替换
    full_text = para.text.replace(old_text, new_text)
    
    # 清空所有 runs 的文本
    for r in runs:
        r.text = ""
    
    # 将替换后的完整文本设置到第一个 run
    runs[0].text = full_text
    return True


def replace_paragraph_text_by_pattern(para, old_pattern, new_text):
    """
    使用正则表达式替换段落中的文本
    
    Args:
        para: 段落对象
        old_pattern: 正则表达式模式
        new_text: 新文本
    """
    import re as re_module
    
    text = para.text
    match = re_module.search(old_pattern, text)
    
    if not match:
        return False
    
    # 替换匹配的文本
    replaced_text = re_module.sub(old_pattern, new_text, text)
    
    runs = para.runs
    if not runs:
        return False
    
    # 清空所有 runs，设置第一个 run
    for r in runs:
        r.text = ""
    runs[0].text = replaced_text
    return True


def add_signature(doc, owners):
    """
    在文档中添加签名信息
    单位显示名称，个人留空
    """
    body = doc.element.body
    paras = body.findall(qn("w:p"))
    start = -1
    found = False

    for i, p in enumerate(paras):
        t = "".join(x.text or "" for x in p.findall(".//" + qn("w:t")))
        if "一式" in t:
            found = True
        if found and t.strip() == "甲方：":
            start = i
            break

    if start < 0:
        return

    names = [
        "甲方：",
        "乙方：",
        "丙方：",
        "丁方：",
        "戊方：",
        "己方：",
        "庚方：",
        "辛方：",
        "壬方：",
        "癸方：",
        "第十一方：",
        "第十二方：",
        "第十三方：",
    ]

    for i, p in enumerate(paras[start:], start=start):
        t = "".join(x.text or "" for x in p.findall(".//" + qn("w:t"))).strip()
        if t in names:
            idx = names.index(t)
            if idx < len(owners) and not owners[idx]["is_person"]:
                runs = p.findall(".//" + qn("w:r"))
                if runs:
                    nr = OxmlElement("w:r")
                    rp = runs[0].find(qn("w:rPr"))
                    if rp is not None:
                        nr.append(deepcopy(rp))
                    nt = OxmlElement("w:t")
                    nt.text = owners[idx]["name"]
                    nr.append(nt)
                    p.append(nr)


def generate_agreement(app_file_path, template_dir, output_dir, custom_agreement_date=None):
    """
    生成合作协议（支持多个软件，生成在一个文档中）

    Args:
        app_file_path: 申请表文件路径
        template_dir: 模板目录
        output_dir: 输出目录
        custom_agreement_date: 自定义协议签署日期（可选）

    Returns:
        dict: {"ok": True, "output": "输出文件路径"} 或 {"ok": False, "err": "错误信息"}
    """
    import re as re_module  # 添加这行导入
    
    # 提取所有软件信息
    software_list = extract_info(app_file_path)

    if not software_list:
        return {"ok": False, "err": "无法提取软件信息"}

    # 检查第一个软件的著作权人
    if not software_list[0].get("owners"):
        return {"ok": False, "err": "无法提取著作权人信息"}

    # 著作权人数量（所有软件共享）
    n = len(software_list[0]["owners"])

    # 选择模板
    template_file = os.path.join(template_dir, f"合作协议-{n}方.docx")
    if not os.path.exists(template_file):
        return {"ok": False, "err": f"模板不存在: {n}方.docx"}

    # 加载模板文档（直接在模板上修改，保留所有格式）
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from copy import deepcopy
    
    # 提取模板中的占位符
    import zipfile
    import xml.etree.ElementTree as ET
    
    template_placeholders = {"software": None, "date": None, "names": []}
    
    with zipfile.ZipFile(template_file) as z:
        xml_content = z.read('word/document.xml')
        root = ET.fromstring(xml_content)
        
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # 提取所有段落文本
        for para in root.findall('.//w:p', ns):
            para_text = ''.join([t.text for t in para.findall('.//w:t', ns) if t.text])
            
            # 查找软件名称占位符（引号内的V1.0等）
            if '合作开发' in para_text and 'V' in para_text:
                # 匹配中文引号（U+201C和U+201D）和英文引号
                match = re_module.search(r'[\u201c\u201d"]([^"""]+V\d+\.\d+[^"""]*)[\u201c\u201d"]', para_text)
                if match:
                    template_placeholders["software"] = match.group(1)
            
            # 查找日期占位符
            if '日期：' in para_text:
                match = re_module.search(r'(\d{4}[-年/]\d{1,2}[-月/]\d{1,2}[日]?)', para_text)
                if match:
                    template_placeholders["date"] = match.group(1)
            
            # 查找示例姓名（2-4个汉字，且不是"甲方"等标题）
            if len(para_text) <= 10 and para_text not in ['甲方：', '乙方：', '丙方：', '丁方：', '戊方：', '己方：', '庚方：', '辛方：', '壬方：', '癸方：', '第十一方：', '第十二方：', '第十三方：']:
                if re_module.match(r'^[\u4e00-\u9fa5]{2,4}$', para_text):
                    template_placeholders["names"].append(para_text)
    
    # 为每个软件生成独立的协议文档
    output_files = []
    
    for idx, software_info in enumerate(software_list):
        # 加载模板文档（每次都重新加载，保留所有格式）
        doc = Document(template_file)
        
        # 准备替换内容
        software_full = f"{software_info['name']}{software_info['version']}"
        agreement_date = calc_agreement_date(software_info["date"], custom_agreement_date)

        # 构建替换列表（使用模板中的实际占位符）
        reps = []
        
        # 1. 软件名称
        if template_placeholders["software"]:
            reps.append((template_placeholders["software"], software_full))
        
        # 2. 日期
        if template_placeholders["date"]:
            reps.append((template_placeholders["date"], agreement_date))
        
        # 3. 示例姓名
        for i, o in enumerate(software_info["owners"]):
            if i < len(template_placeholders["names"]):
                reps.append((template_placeholders["names"][i], o["name"]))
        
        # 4. 份数
        for nc in ["十五", "十四", "十三", "十二", "十一", "十", "九", "八", "七", "六", "五", "四", "三"]:
            reps.append((f"{nc}份", f"{n+1}份"))
        
        # 执行替换（遍历段落和表格）
        for old, new in reps:
            for para in doc.paragraphs:
                replace_paragraph_text(para, old, new)
            for tbl in doc.tables:
                for r in tbl.rows:
                    for c in r.cells:
                        for para in c.paragraphs:
                            replace_paragraph_text(para, old, new)

        # 5. 添加签名
        add_signature(doc, software_info["owners"])

        # 保存文件（使用当前软件的名称）
        safe_name = software_info["name"].replace("/", "-").replace("\\", "-").replace(":", "-")
        output_file = os.path.join(output_dir, f"合作协议-{safe_name}.docx")
        doc.save(output_file)
        output_files.append(output_file)

    # 返回结果
    # 注意：API只返回第一个文件的下载链接（兼容现有前端）
    # 如果需要支持多文件下载，需要修改API和前端
    return {
        "ok": True, 
        "output": output_files[0] if output_files else "",
        "output_files": output_files,  # 所有生成的文件
        "software_count": len(software_list),
        "software_list": [s["name"] for s in software_list]
    }
