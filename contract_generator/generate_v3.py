# -*- coding: utf-8 -*-
"""
合作协议批量生成器 v3 - 改进文本替换
"""

import os
import sys
import glob
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from copy import deepcopy

# Windows Python 路径
sys.path.insert(0, r'C:\Users\yuran\AppData\Local\Programs\Python\Python37\Lib\site-packages')

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)


TIAN_GAN = ['甲', '乙', '丙', '丁', '戊', '己', '庚', '辛', '壬', '癸']


def get_party_title(index: int) -> str:
    if index < 10:
        return f'{TIAN_GAN[index]}方'
    else:
        return f'第{index + 1}方'


def parse_date(date_str: str) -> datetime:
    date_str = date_str.strip()
    if '年' in date_str and '月' in date_str and '日' in date_str:
        try:
            parts = date_str.replace('年', ' ').replace('月', ' ').replace('日', '').split()
            return datetime(int(parts[0]), int(parts[1]), int(parts[2]))
        except:
            pass
    return datetime.now()


def calc_contract_date(completion_str: str) -> str:
    dt = parse_date(completion_str)
    y, m = dt.year, dt.month - 3
    if m <= 0:
        m += 12
        y -= 1
    return f'{y}年{m}月{dt.day}日'


def parse_docx_tables(docx_path: str) -> list:
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    tables = []
    with zipfile.ZipFile(docx_path, 'r') as zf:
        xml_content = zf.read('word/document.xml')
        tree = ET.fromstring(xml_content)
        for tbl in tree.findall('.//w:tbl', ns):
            table_data = []
            for row in tbl.findall('.//w:tr', ns):
                row_data = []
                for cell in row.findall('.//w:tc', ns):
                    cell_text = ''
                    for text in cell.findall('.//w:t', ns):
                        if text.text:
                            cell_text += text.text
                    row_data.append(cell_text.strip())
                if row_data:
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)
    return tables


def is_individual(id_number: str) -> bool:
    """判断是否为自然人（个人）"""
    if not id_number:
        return False
    # 身份证号：18位数字（最后一位可能是X）
    import re
    return bool(re.match(r'^\d{17}[\dXx]$', id_number.strip()))


def extract_info(tables: list) -> dict:
    info = {
        'software_name': '',
        'version': 'V1.0',
        'completion_date': '',
        'owners': []
    }
    for table in tables:
        for i in range(len(table)):
            row = table[i]
            row_text = ' '.join(row)
            if '软著名称' in row_text and i + 1 < len(table):
                data = table[i + 1]
                info['software_name'] = data[0] if data else ''
                info['version'] = data[2] if len(data) > 2 and data[2] else 'V1.0'
                info['completion_date'] = data[3] if len(data) > 3 else ''
            if '著作权人信息' in row_text:
                for j in range(i + 1, min(i + 3, len(table))):
                    if '公司/单位' in ' '.join(table[j]):
                        for k in range(j + 1, len(table)):
                            r = table[k]
                            if '联系人' in ' '.join(r) or not r or not r[0]:
                                break
                            if '公司' in r[0] or '营业执照' in r[0]:
                                continue
                            id_num = r[1] if len(r) > 1 else ''
                            info['owners'].append({
                                'name': r[0],
                                'id_number': id_num,
                                'location': r[2] if len(r) > 2 else '',
                                'is_individual': is_individual(id_num)  # 是否为个人
                            })
                        break
    return info


def replace_in_runs(paragraph, old_text: str, new_text: str):
    """改进的文本替换：处理 run 级别的拆分"""
    full_text = paragraph.text
    
    if old_text not in full_text:
        return False
    
    runs = paragraph.runs
    if not runs:
        return False
    
    first_run = runs[0]
    new_full_text = full_text.replace(old_text, new_text)
    
    for run in runs:
        run.text = ''
    
    first_run.text = new_full_text
    return True


def add_signature_names(doc, owners: list):
    """在签名部分添加名字（单位写名，个人留空）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    body = doc.element.body
    paras = body.findall(qn('w:p'))
    
    # 找到签名部分的"甲方："位置（在"一式"之后）
    signature_start = -1
    found_yishi = False
    for i, para in enumerate(paras):
        text = ''.join(t.text for t in para.findall('.//' + qn('w:t')) if t.text)
        text = text.strip()
        
        if '一式' in text:
            found_yishi = True
        
        if found_yishi and text == '甲方：':
            signature_start = i
            break
    
    if signature_start < 0:
        return
    
    # 处理每个"XX方："段落
    party_names = ['甲方：', '乙方：', '丙方：', '丁方：', '戊方：', '己方：', 
                   '庚方：', '辛方：', '壬方：', '癸方：', '第十一方：', '第十二方：', '第十三方：']
    
    for i, para in enumerate(paras[signature_start:], start=signature_start):
        text = ''.join(t.text for t in para.findall('.//' + qn('w:t')) if t.text)
        text = text.strip()
        
        if text in party_names:
            party_idx = party_names.index(text)
            if party_idx < len(owners):
                owner = owners[party_idx]
                
                if not owner.get('is_individual'):
                    # 单位：在"甲方："后面追加名字
                    runs = para.findall('.//' + qn('w:r'))
                    if runs:
                        first_run = runs[0]
                        new_run = OxmlElement('w:r')
                        rPr = first_run.find(qn('w:rPr'))
                        if rPr is not None:
                            new_run.append(deepcopy(rPr))
                        new_text = OxmlElement('w:t')
                        new_text.text = owner['name']
                        new_run.append(new_text)
                        para.append(new_run)
                # 个人：保持原样（"乙方："不动，留空用于手写签名）


def replace_in_table(doc, old_text: str, new_text: str):
    """替换表格中的文本"""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, old_text, new_text)


def generate_contract(app_path: str, template_dir: str, output_dir: str) -> dict:
    # 提取信息
    tables = parse_docx_tables(app_path)
    info = extract_info(tables)
    
    if not info['software_name']:
        return {'success': False, 'error': '无法提取软件名称'}
    if not info['owners']:
        return {'success': False, 'error': '无法提取著作权人'}
    
    owner_count = len(info['owners'])
    template_path = os.path.join(template_dir, f'合作协议-{owner_count}方.docx')
    
    if not os.path.exists(template_path):
        return {'success': False, 'error': f'模板不存在: {owner_count}方'}
    
    contract_date = calc_contract_date(info['completion_date'])
    software_full = f"{info['software_name']}{info['version']}"
    copy_count = owner_count + 1
    
    # 打开模板
    doc = Document(template_path)
    
    # 要替换的内容
    replacements = [
        ('智能心理压力情绪疏导与调解平台V1.0', software_full),
        ('2025-3-10', contract_date),
    ]
    
    # 添加著作权人替换（开头部分，无论个人单位都写名字）
    sample_names = ['王红梅', '汪庆军', '吴六韬', '赖羽羽', '刘茜茜',
                    '丁姿慧', '刘澈', '曹冬梅', '谢奇秀', '钟海玲']
    for i, owner in enumerate(info['owners']):
        if i < len(sample_names):
            replacements.append((sample_names[i], owner['name']))
    
    # 份数替换
    num_chars = ['三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二', '十三', '十四']
    for nc in num_chars:
        replacements.append((f'{nc}份', f'{copy_count}份'))
    
    # 执行替换
    for old_text, new_text in replacements:
        # 替换段落
        for para in doc.paragraphs:
            replace_in_runs(para, old_text, new_text)
        # 替换表格
        replace_in_table(doc, old_text, new_text)
    
    # 在签名部分添加名字（单位写名，个人留空）
    add_signature_names(doc, info['owners'])
    
    # 保存
    safe_name = info['software_name'].replace('/', '-').replace('\\', '-').replace(':', '-')
    output_path = os.path.join(output_dir, f'合作协议-{safe_name}.docx')
    doc.save(output_path)
    
    return {
        'success': True,
        'software_name': info['software_name'],
        'owners_count': owner_count,
        'owners': [{'name': o['name'], 'is_individual': o.get('is_individual', False)} for o in info['owners']],
        'output_file': output_path
    }


def batch_process(input_dir: str, output_dir: str, template_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    files = glob.glob(os.path.join(input_dir, '*.docx'))
    
    if not files:
        print(f"没有找到申请表文件: {input_dir}")
        return
    
    print(f"找到 {len(files)} 个申请表文件")
    
    results = []
    for f in files:
        print(f"\n处理: {os.path.basename(f)}")
        
        try:
            result = generate_contract(f, template_dir, output_dir)
            results.append(result)
            
            if result['success']:
                print(f"  软件: {result['software_name']}")
                owners_info = []
                for o in result['owners']:
                    if o['is_individual']:
                        owners_info.append(f"{o['name']}(个人)")
                    else:
                        owners_info.append(f"{o['name']}(单位)")
                print(f"  著作权人: {', '.join(owners_info)}")
                print(f"  [OK] 生成成功")
            else:
                print(f"  [FAIL] {result['error']}")
        except Exception as e:
            print(f"  [ERROR] {e}")
            import traceback
            traceback.print_exc()
            results.append({'success': False, 'error': str(e)})
    
    success = sum(1 for r in results if r.get('success'))
    print(f"\n处理完成! 成功: {success}/{len(results)}")
    
    return results


if __name__ == '__main__':
    print("=" * 50)
    print("  软件著作权合作协议批量生成器 v3")
    print("=" * 50)
    
    input_dir = r'C:\Temp\contracts\input'
    output_dir = r'C:\Temp\contracts\output'
    template_dir = r'C:\Temp\contracts\中安-合作协议'
    
    print(f"\n输入: {input_dir}")
    print(f"输出: {output_dir}")
    print(f"模板: {template_dir}")
    
    batch_process(input_dir, output_dir, template_dir)
