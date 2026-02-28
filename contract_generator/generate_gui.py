# -*- coding: utf-8 -*-
"""
软件著作权合作协议生成器 - GUI 版本
"""

import os
import sys
import shutil
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from tkinter import *
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)

from copy import deepcopy

TIAN_GAN = ['甲', '乙', '丙', '丁', '戊', '己', '庚', '辛', '壬', '癸']


class ContractGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("软件著作权合作协议生成器")
        self.root.geometry("800x600")
        self.root.resizable(True, True)
        
        # 默认路径
        self.desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        if not os.path.exists(self.desktop):
            self.desktop = os.path.join(os.path.expanduser('~'), '桌面')
        
        self.template_dir = r"C:\Temp\contracts\中安-合作协议"
        self.output_dir = self.desktop
        
        self.input_files = []
        
        self.setup_ui()
        
    def setup_ui(self):
        """设置界面"""
        # 标题
        title_frame = Frame(self.root, bg="#2c3e50", height=60)
        title_frame.pack(fill=X)
        title_frame.pack_propagate(False)
        
        title_label = Label(title_frame, text="软件著作权合作协议生成器", 
                           font=("Microsoft YaHei", 16, "bold"), 
                           bg="#2c3e50", fg="white")
        title_label.pack(pady=15)
        
        # 主区域
        main_frame = Frame(self.root, padx=20, pady=20)
        main_frame.pack(fill=BOTH, expand=True)
        
        # 文件拖放区
        drop_frame = LabelFrame(main_frame, text=" 申请表文件 ", 
                               font=("Microsoft YaHei", 10), padx=10, pady=10)
        drop_frame.pack(fill=X, pady=(0, 10))
        
        # 文件列表
        self.file_frame = Frame(drop_frame)
        self.file_frame.pack(fill=X)
        
        self.file_label = Label(self.file_frame, text="将 .docx 文件拖到此处或点击选择", 
                               font=("Microsoft YaHei", 9), fg="#7f8c8d", height=3)
        self.file_label.pack(side=LEFT, padx=10)
        
        btn_frame = Frame(self.file_frame)
        btn_frame.pack(side=RIGHT)
        
        self.select_btn = Button(btn_frame, text="选择文件", command=self.select_files,
                                font=("Microsoft YaHei", 9), width=10)
        self.select_btn.pack(side=LEFT, padx=5)
        
        self.clear_btn = Button(btn_frame, text="清空", command=self.clear_files,
                               font=("Microsoft YaHei", 9), width=8)
        self.clear_btn.pack(side=LEFT)
        
        # 文件列表框
        self.file_listbox = Listbox(drop_frame, height=4, font=("Microsoft YaHei", 9))
        self.file_listbox.pack(fill=X, pady=(10, 0))
        
        # 绑定拖放事件
        self.file_listbox.bind('<Double-Button-1>', self.preview_file)
        
        # 设置区
        settings_frame = LabelFrame(main_frame, text=" 设置 ", 
                                   font=("Microsoft YaHei", 10), padx=10, pady=10)
        settings_frame.pack(fill=X, pady=(0, 10))
        
        # 模板路径
        row1 = Frame(settings_frame)
        row1.pack(fill=X, pady=5)
        
        Label(row1, text="模板目录：", font=("Microsoft YaHei", 9), width=10, anchor='w').pack(side=LEFT)
        self.template_entry = Entry(row1, font=("Microsoft YaHei", 9))
        self.template_entry.insert(0, self.template_dir)
        self.template_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
        Button(row1, text="浏览", command=self.select_template_dir, width=6).pack(side=LEFT)
        
        # 输出路径
        row2 = Frame(settings_frame)
        row2.pack(fill=X, pady=5)
        
        Label(row2, text="输出目录：", font=("Microsoft YaHei", 9), width=10, anchor='w').pack(side=LEFT)
        self.output_entry = Entry(row2, font=("Microsoft YaHei", 9))
        self.output_entry.insert(0, self.output_dir)
        self.output_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
        Button(row2, text="浏览", command=self.select_output_dir, width=6).pack(side=LEFT)
        
        # 预览区
        preview_frame = LabelFrame(main_frame, text=" 预览 ", 
                                  font=("Microsoft YaHei", 10), padx=10, pady=10)
        preview_frame.pack(fill=BOTH, expand=True, pady=(0, 10))
        
        self.preview_text = ScrolledText(preview_frame, height=8, font=("Microsoft YaHei", 9))
        self.preview_text.pack(fill=BOTH, expand=True)
        self.preview_text.insert('1.0', "双击文件列表中的文件可预览...")
        
        # 按钮区
        btn_frame = Frame(main_frame)
        btn_frame.pack(fill=X)
        
        self.preview_btn = Button(btn_frame, text="预览", command=self.preview_selected,
                                 font=("Microsoft YaHei", 10), width=12, bg="#3498db", fg="white")
        self.preview_btn.pack(side=LEFT, padx=5)
        
        self.generate_btn = Button(btn_frame, text="生成协议", command=self.generate,
                                  font=("Microsoft YaHei", 10, "bold"), width=15, bg="#27ae60", fg="white")
        self.generate_btn.pack(side=LEFT, padx=5)
        
        self.open_btn = Button(btn_frame, text="打开输出目录", command=self.open_output_dir,
                              font=("Microsoft YaHei", 10), width=12)
        self.open_btn.pack(side=LEFT, padx=5)
        
        # 状态栏
        self.status_var = StringVar(value="就绪")
        status_bar = Label(self.root, textvariable=self.status_var, 
                          font=("Microsoft YaHei", 9), bg="#ecf0f1", anchor='w', padx=10)
        status_bar.pack(fill=X, side=BOTTOM)
        
        # 绑定拖放
        self.setup_drag_drop()
        
    def setup_drag_drop(self):
        """设置拖放功能"""
        # tkinter 原生不支持拖放，使用点击选择代替
        pass
        
    def select_files(self):
        """选择文件"""
        files = filedialog.askopenfilenames(
            title="选择申请表文件",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if files:
            for f in files:
                if f not in self.input_files:
                    self.input_files.append(f)
                    self.file_listbox.insert(END, os.path.basename(f))
            self.status_var.set(f"已选择 {len(self.input_files)} 个文件")
            
    def clear_files(self):
        """清空文件列表"""
        self.input_files.clear()
        self.file_listbox.delete(0, END)
        self.preview_text.delete('1.0', END)
        self.preview_text.insert('1.0', "双击文件列表中的文件可预览...")
        self.status_var.set("已清空")
        
    def select_template_dir(self):
        """选择模板目录"""
        d = filedialog.askdirectory(initialdir=self.template_dir, title="选择模板目录")
        if d:
            self.template_dir = d
            self.template_entry.delete(0, END)
            self.template_entry.insert(0, d)
            
    def select_output_dir(self):
        """选择输出目录"""
        d = filedialog.askdirectory(initialdir=self.output_dir, title="选择输出目录")
        if d:
            self.output_dir = d
            self.output_entry.delete(0, END)
            self.output_entry.insert(0, d)
            
    def open_output_dir(self):
        """打开输出目录"""
        output_dir = self.output_entry.get()
        if os.path.exists(output_dir):
            os.startfile(output_dir)
        else:
            messagebox.showwarning("提示", "输出目录不存在")
            
    def preview_file(self, event):
        """双击预览文件"""
        selection = self.file_listbox.curselection()
        if selection:
            self.preview_selected()
            
    def preview_selected(self):
        """预览选中的文件"""
        selection = self.file_listbox.curselection()
        if not selection:
            messagebox.showinfo("提示", "请先选择要预览的文件")
            return
            
        idx = selection[0]
        file_path = self.input_files[idx]
        
        try:
            info = self.extract_info_from_docx(file_path)
            
            preview = f"""【文件信息】
文件名: {os.path.basename(file_path)}

【提取结果】
软件名称: {info['software_name']}
版本号: {info['version']}
开发完成日期: {info['completion_date']}
协议日期: {self.calc_contract_date(info['completion_date'])}

【著作权人】({len(info['owners'])}人)
"""
            for i, owner in enumerate(info['owners']):
                owner_type = "个人" if owner.get('is_individual') else "单位"
                preview += f"  {i+1}. {owner['name']} ({owner_type})\n"
            
            preview += f"""
【将使用模板】
合作协议-{len(info['owners'])}方.docx

【协议份数】
{len(info['owners']) + 1} 份
"""
            
            self.preview_text.delete('1.0', END)
            self.preview_text.insert('1.0', preview)
            
        except Exception as e:
            self.preview_text.delete('1.0', END)
            self.preview_text.insert('1.0', f"预览失败: {str(e)}")
            
    def extract_info_from_docx(self, docx_path):
        """从 docx 提取信息"""
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        tables = []
        with zipfile.ZipFile(docx_path, 'r') as zf:
            xml_content = zf.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            
            for tbl in tree.findall('.//w:tbl', ns):
                table_data = []
                for row in tbl.findall('.//w:tr', ns):
                    row_data = []
                    for cell in row.findall('.//w:tc', ns):
                        cell_text = ''
                        for text in cell.findall('.//w:t', ns):
                            if text.text:
                                cell_text += text.text
                        row_data.append(cell_text.strip())
                    if row_data:
                        table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
        
        info = {
            'software_name': '',
            'version': 'V1.0',
            'completion_date': '',
            'owners': []
        }
        
        for table in tables:
            for i in range(len(table)):
                row = table[i]
                row_text = ' '.join(row)
                
                if '软著名称' in row_text and i + 1 < len(table):
                    data = table[i + 1]
                    info['software_name'] = data[0] if data else ''
                    info['version'] = data[2] if len(data) > 2 and data[2] else 'V1.0'
                    info['completion_date'] = data[3] if len(data) > 3 else ''
                
                if '著作权人信息' in row_text:
                    for j in range(i + 1, min(i + 3, len(table))):
                        if '公司/单位' in ' '.join(table[j]):
                            for k in range(j + 1, len(table)):
                                r = table[k]
                                if '联系人' in ' '.join(r) or not r or not r[0]:
                                    break
                                if '公司' in r[0] or '营业执照' in r[0]:
                                    continue
                                id_num = r[1] if len(r) > 1 else ''
                                info['owners'].append({
                                    'name': r[0],
                                    'id_number': id_num,
                                    'location': r[2] if len(r) > 2 else '',
                                    'is_individual': self.is_individual(id_num)
                                })
                            break
        return info
        
    def is_individual(self, id_number):
        """判断是否为个人"""
        if not id_number:
            return False
        import re
        return bool(re.match(r'^\d{17}[\dXx]$', id_number.strip()))
        
    def calc_contract_date(self, completion_str):
        """计算协议日期"""
        try:
            if '年' in completion_str:
                parts = completion_str.replace('年', ' ').replace('月', ' ').replace('日', '').split()
                y, m, d = int(parts[0]), int(parts[1]), int(parts[2])
                m = m - 3
                if m <= 0:
                    m += 12
                    y -= 1
                return f'{y}年{m}月{d}日'
        except:
            pass
        return '2025年1月1日'
        
    def generate(self):
        """生成协议"""
        if not self.input_files:
            messagebox.showwarning("提示", "请先选择申请表文件")
            return
            
        output_dir = self.output_entry.get()
        template_dir = self.template_entry.get()
        
        if not os.path.exists(template_dir):
            messagebox.showerror("错误", f"模板目录不存在: {template_dir}")
            return
            
        os.makedirs(output_dir, exist_ok=True)
        
        success_count = 0
        fail_count = 0
        results = []
        
        for file_path in self.input_files:
            try:
                result = self.generate_contract(file_path, template_dir, output_dir)
                if result['success']:
                    success_count += 1
                    results.append(f"✓ {os.path.basename(file_path)}\n  → {os.path.basename(result['output_file'])}")
                else:
                    fail_count += 1
                    results.append(f"✗ {os.path.basename(file_path)}\n  错误: {result['error']}")
            except Exception as e:
                fail_count += 1
                results.append(f"✗ {os.path.basename(file_path)}\n  错误: {str(e)}")
        
        # 显示结果
        msg = f"处理完成！\n\n成功: {success_count} 个\n失败: {fail_count} 个\n\n"
        msg += "\n".join(results)
        
        if success_count > 0:
            messagebox.showinfo("完成", msg)
        else:
            messagebox.showwarning("完成", msg)
            
        self.status_var.set(f"成功: {success_count}, 失败: {fail_count}")
        
    def generate_contract(self, app_path, template_dir, output_dir):
        """生成单个协议"""
        info = self.extract_info_from_docx(app_path)
        
        if not info['software_name']:
            return {'success': False, 'error': '无法提取软件名称'}
        if not info['owners']:
            return {'success': False, 'error': '无法提取著作权人'}
        
        owner_count = len(info['owners'])
        template_path = os.path.join(template_dir, f'合作协议-{owner_count}方.docx')
        
        if not os.path.exists(template_path):
            return {'success': False, 'error': f'模板不存在: {owner_count}方.docx'}
        
        contract_date = self.calc_contract_date(info['completion_date'])
        software_full = f"{info['software_name']}{info['version']}"
        copy_count = owner_count + 1
        
        doc = Document(template_path)
        
        # 替换文本
        replacements = [
            ('智能心理压力情绪疏导与调解平台V1.0', software_full),
            ('2025-3-10', contract_date),
        ]
        
        sample_names = ['王红梅', '汪庆军', '吴六韬', '赖羽羽', '刘茜茜',
                        '丁姿慧', '刘澈', '曹冬梅', '谢奇秀', '钟海玲']
        for i, owner in enumerate(info['owners']):
            if i < len(sample_names):
                replacements.append((sample_names[i], owner['name']))
        
        num_chars = ['三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二', '十三', '十四']
        for nc in num_chars:
            replacements.append((f'{nc}份', f'{copy_count}份'))
        
        for old_text, new_text in replacements:
            for para in doc.paragraphs:
                self.replace_in_runs(para, old_text, new_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self.replace_in_runs(para, old_text, new_text)
        
        # 添加签名
        self.add_signature_names(doc, info['owners'])
        
        # 保存
        safe_name = info['software_name'].replace('/', '-').replace('\\', '-').replace(':', '-')
        output_path = os.path.join(output_dir, f'合作协议-{safe_name}.docx')
        doc.save(output_path)
        
        return {
            'success': True,
            'software_name': info['software_name'],
            'owners_count': owner_count,
            'output_file': output_path
        }
        
    def replace_in_runs(self, paragraph, old_text, new_text):
        """替换段落中的文本"""
        if old_text not in paragraph.text:
            return
        runs = paragraph.runs
        if not runs:
            return
        full_text = paragraph.text
        new_full_text = full_text.replace(old_text, new_text)
        for run in runs:
            run.text = ''
        runs[0].text = new_full_text
        
    def add_signature_names(self, doc, owners):
        """在签名部分添加名字"""
        body = doc.element.body
        paras = body.findall(qn('w:p'))
        
        signature_start = -1
        found_yishi = False
        for i, para in enumerate(paras):
            text = ''.join(t.text for t in para.findall('.//' + qn('w:t')) if t.text)
            text = text.strip()
            
            if '一式' in text:
                found_yishi = True
            
            if found_yishi and text == '甲方：':
                signature_start = i
                break
        
        if signature_start < 0:
            return
        
        party_names = ['甲方：', '乙方：', '丙方：', '丁方：', '戊方：', '己方：', 
                       '庚方：', '辛方：', '壬方：', '癸方：', '第十一方：', '第十二方：', '第十三方：']
        
        for i, para in enumerate(paras[signature_start:], start=signature_start):
            text = ''.join(t.text for t in para.findall('.//' + qn('w:t')) if t.text)
            text = text.strip()
            
            if text in party_names:
                party_idx = party_names.index(text)
                if party_idx < len(owners):
                    owner = owners[party_idx]
                    
                    if not owner.get('is_individual'):
                        runs = para.findall('.//' + qn('w:r'))
                        if runs:
                            first_run = runs[0]
                            new_run = OxmlElement('w:r')
                            rPr = first_run.find(qn('w:rPr'))
                            if rPr is not None:
                                new_run.append(deepcopy(rPr))
                            new_text = OxmlElement('w:t')
                            new_text.text = owner['name']
                            new_run.append(new_text)
                            para.append(new_run)
                            
    def convert_doc_to_docx(self, doc_path):
        """转换 .doc 为 .docx"""
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(doc_path)
            docx_path = doc_path.replace('.doc', '.docx')
            doc.SaveAs(docx_path, 16)
            doc.Close()
            word.Quit()
            return docx_path
        except:
            return None


def main():
    root = Tk()
    app = ContractGenerator(root)
    root.mainloop()


if __name__ == '__main__':
    main()
