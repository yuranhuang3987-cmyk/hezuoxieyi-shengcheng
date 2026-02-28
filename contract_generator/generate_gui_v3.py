# -*- coding: utf-8 -*-
"""
软件著作权合作协议生成器 - GUI 版本 v3
支持 Windows 拖放 (windnd)
"""

import os
import sys
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from tkinter import *
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)

try:
    import windnd
    DRAG_DROP_AVAILABLE = True
except ImportError:
    DRAG_DROP_AVAILABLE = False
    print("提示: 安装 windnd 可支持拖放功能: pip install windnd")

from copy import deepcopy


class ContractGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("软件著作权合作协议生成器")
        self.root.geometry("800x700")
        self.root.minsize(700, 600)
        
        # 默认路径
        self.desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
        if not os.path.exists(self.desktop):
            self.desktop = os.path.join(os.path.expanduser('~'), '桌面')
        
        self.template_dir = r"C:\Temp\contracts\中安-合作协议"
        self.output_dir = self.desktop
        self.input_files = []
        
        self.setup_ui()
        
        # 设置拖放
        if DRAG_DROP_AVAILABLE:
            windnd.hook_dropfiles(self.root, func=self.on_drop)
        
    def on_drop(self, files):
        """处理拖放"""
        for f in files:
            f = f.decode('utf-8') if isinstance(f, bytes) else f
            if os.path.isfile(f) and f.lower().endswith('.docx'):
                if f not in self.input_files:
                    self.input_files.append(f)
                    self.file_listbox.insert(END, os.path.basename(f))
            elif os.path.isdir(f):
                for item in os.listdir(f):
                    if item.lower().endswith('.docx'):
                        full_path = os.path.join(f, item)
                        if full_path not in self.input_files:
                            self.input_files.append(full_path)
                            self.file_listbox.insert(END, item)
        
        self.count_label.config(text=f"已选择: {len(self.input_files)} 个文件")
        self.status_var.set(f"已添加 {len(self.input_files)} 个文件")
        
    def setup_ui(self):
        """设置界面"""
        # 主框架
        main_frame = Frame(self.root, padx=20, pady=20)
        main_frame.pack(fill=BOTH, expand=True)
        
        # 标题
        title = Label(main_frame, text="软件著作权合作协议生成器", 
                     font=("Microsoft YaHei", 18, "bold"), fg="#2c3e50")
        title.pack(pady=(0, 20))
        
        # 文件选择区
        file_frame = LabelFrame(main_frame, text=" 选择申请表 (可拖动文件到此处) ", 
                               font=("Microsoft YaHei", 10), padx=10, pady=10)
        file_frame.pack(fill=X, pady=(0, 10))
        
        # 按钮行
        btn_row = Frame(file_frame)
        btn_row.pack(fill=X, pady=5)
        
        self.add_btn = Button(btn_row, text="添加文件", command=self.add_files,
                             font=("Microsoft YaHei", 10), width=12, bg="#3498db", fg="white")
        self.add_btn.pack(side=LEFT, padx=5)
        
        self.add_folder_btn = Button(btn_row, text="添加文件夹", command=self.add_folder,
                                    font=("Microsoft YaHei", 10), width=12)
        self.add_folder_btn.pack(side=LEFT, padx=5)
        
        self.clear_btn = Button(btn_row, text="清空列表", command=self.clear_files,
                               font=("Microsoft YaHei", 10), width=12)
        self.clear_btn.pack(side=LEFT, padx=5)
        
        # 文件数量标签
        self.count_label = Label(btn_row, text="已选择: 0 个文件", 
                                font=("Microsoft YaHei", 9), fg="#7f8c8d")
        self.count_label.pack(side=RIGHT, padx=10)
        
        # 文件列表
        list_frame = Frame(file_frame)
        list_frame.pack(fill=X, pady=5)
        
        self.file_listbox = Listbox(list_frame, height=5, font=("Microsoft YaHei", 9),
                                   selectmode=EXTENDED)
        scrollbar = Scrollbar(list_frame, orient=VERTICAL, command=self.file_listbox.yview)
        self.file_listbox.configure(yscrollcommand=scrollbar.set)
        
        self.file_listbox.pack(side=LEFT, fill=X, expand=True)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        self.file_listbox.bind('<Double-Button-1>', self.preview_file)
        self.file_listbox.bind('<Delete>', self.remove_selected)
        
        # 提示
        if DRAG_DROP_AVAILABLE:
            tip = Label(file_frame, text="提示: 直接拖动 .docx 文件到窗口 | 双击预览 | Delete删除", 
                       font=("Microsoft YaHei", 8), fg="#27ae60")
        else:
            tip = Label(file_frame, text="提示: 双击文件预览 | 选中后按 Delete 删除", 
                       font=("Microsoft YaHei", 8), fg="#95a5a6")
        tip.pack(anchor='w')
        
        # 设置区
        settings_frame = LabelFrame(main_frame, text=" 设置 ", 
                                   font=("Microsoft YaHei", 10), padx=10, pady=10)
        settings_frame.pack(fill=X, pady=(0, 10))
        
        # 输出目录
        row1 = Frame(settings_frame)
        row1.pack(fill=X, pady=5)
        
        Label(row1, text="输出到:", font=("Microsoft YaHei", 9), width=8, anchor='w').pack(side=LEFT)
        self.output_entry = Entry(row1, font=("Microsoft YaHei", 9))
        self.output_entry.insert(0, self.output_dir)
        self.output_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
        Button(row1, text="浏览", command=self.select_output_dir, width=6).pack(side=LEFT)
        
        # 模板目录
        row2 = Frame(settings_frame)
        row2.pack(fill=X, pady=5)
        
        Label(row2, text="模板:", font=("Microsoft YaHei", 9), width=8, anchor='w').pack(side=LEFT)
        self.template_entry = Entry(row2, font=("Microsoft YaHei", 9))
        self.template_entry.insert(0, self.template_dir)
        self.template_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
        Button(row2, text="浏览", command=self.select_template_dir, width=6).pack(side=LEFT)
        
        # 预览区
        preview_frame = LabelFrame(main_frame, text=" 预览 ", 
                                  font=("Microsoft YaHei", 10), padx=10, pady=10)
        preview_frame.pack(fill=BOTH, expand=True, pady=(0, 10))
        
        self.preview_text = ScrolledText(preview_frame, height=10, font=("Microsoft YaHei", 9))
        self.preview_text.pack(fill=BOTH, expand=True)
        self.preview_text.insert('1.0', "双击上方文件列表中的文件可预览提取的信息...\n\n或者直接把 .docx 文件拖到这个窗口上！")
        
        # 底部按钮
        bottom_frame = Frame(main_frame)
        bottom_frame.pack(fill=X)
        
        self.preview_btn = Button(bottom_frame, text="预览选中", command=self.preview_selected,
                                 font=("Microsoft YaHei", 10), width=12)
        self.preview_btn.pack(side=LEFT, padx=5)
        
        self.generate_btn = Button(bottom_frame, text="生成协议", command=self.generate,
                                  font=("Microsoft YaHei", 11, "bold"), width=15, 
                                  bg="#27ae60", fg="white")
        self.generate_btn.pack(side=LEFT, padx=10)
        
        self.open_btn = Button(bottom_frame, text="打开输出目录", command=self.open_output_dir,
                              font=("Microsoft YaHei", 10), width=12)
        self.open_btn.pack(side=LEFT, padx=5)
        
        # 状态栏
        self.status_var = StringVar(value="就绪 - 请添加申请表文件（可拖动）")
        status_bar = Label(self.root, textvariable=self.status_var, 
                          font=("Microsoft YaHei", 9), bg="#ecf0f1", anchor='w', padx=10, pady=5)
        status_bar.pack(fill=X, side=BOTTOM)
        
    def add_files(self):
        """添加文件"""
        files = filedialog.askopenfilenames(
            title="选择申请表文件",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        self._add_files_to_list(files)
        
    def add_folder(self):
        """添加文件夹中的所有 .docx"""
        folder = filedialog.askdirectory(title="选择包含申请表的文件夹")
        if folder:
            files = [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith('.docx')]
            self._add_files_to_list(files)
            
    def _add_files_to_list(self, files):
        """添加文件到列表"""
        count = 0
        for f in files:
            if f not in self.input_files:
                self.input_files.append(f)
                self.file_listbox.insert(END, os.path.basename(f))
                count += 1
        
        self.count_label.config(text=f"已选择: {len(self.input_files)} 个文件")
        if count > 0:
            self.status_var.set(f"已添加 {count} 个文件")
            
    def remove_selected(self, event=None):
        """删除选中的文件"""
        selection = self.file_listbox.curselection()
        if selection:
            for i in reversed(selection):
                self.file_listbox.delete(i)
                del self.input_files[i]
            self.count_label.config(text=f"已选择: {len(self.input_files)} 个文件")
            
    def clear_files(self):
        """清空文件列表"""
        self.input_files.clear()
        self.file_listbox.delete(0, END)
        self.preview_text.delete('1.0', END)
        self.preview_text.insert('1.0', "双击上方文件列表中的文件可预览提取的信息...\n\n或者直接把 .docx 文件拖到这个窗口上！")
        self.count_label.config(text="已选择: 0 个文件")
        self.status_var.set("已清空")
        
    def select_output_dir(self):
        """选择输出目录"""
        d = filedialog.askdirectory(initialdir=self.output_entry.get(), title="选择输出目录")
        if d:
            self.output_entry.delete(0, END)
            self.output_entry.insert(0, d)
            
    def select_template_dir(self):
        """选择模板目录"""
        d = filedialog.askdirectory(initialdir=self.template_entry.get(), title="选择模板目录")
        if d:
            self.template_entry.delete(0, END)
            self.template_entry.insert(0, d)
            
    def open_output_dir(self):
        """打开输出目录"""
        output_dir = self.output_entry.get()
        if os.path.exists(output_dir):
            os.startfile(output_dir)
        else:
            os.makedirs(output_dir, exist_ok=True)
            os.startfile(output_dir)
            
    def preview_file(self, event):
        """双击预览"""
        self.preview_selected()
        
    def preview_selected(self):
        """预览选中的文件"""
        selection = self.file_listbox.curselection()
        if not selection:
            messagebox.showinfo("提示", "请先选择要预览的文件")
            return
            
        idx = selection[0]
        file_path = self.input_files[idx]
        
        try:
            info = self.extract_info_from_docx(file_path)
            
            preview = f"""【文件】{os.path.basename(file_path)}

【软件信息】
  名称: {info['software_name']}
  版本: {info['version']}
  开发完成: {info['completion_date']}
  协议日期: {self.calc_contract_date(info['completion_date'])}

【著作权人】共 {len(info['owners'])} 人
"""
            for i, owner in enumerate(info['owners']):
                owner_type = "个人(签名留空)" if owner.get('is_individual') else "单位(显示名称)"
                preview += f"  {i+1}. {owner['name']} [{owner_type}]\n"
            
            template_file = f"合作协议-{len(info['owners'])}方.docx"
            template_path = os.path.join(self.template_entry.get(), template_file)
            
            if os.path.exists(template_path):
                preview += f"\n【模板】{template_file} OK"
            else:
                preview += f"\n【模板】{template_file} 不存在!"
            
            preview += f"\n\n【协议份数】{len(info['owners']) + 1} 份"
            
            self.preview_text.delete('1.0', END)
            self.preview_text.insert('1.0', preview)
            
        except Exception as e:
            self.preview_text.delete('1.0', END)
            self.preview_text.insert('1.0', f"预览失败: {str(e)}")
            
    def extract_info_from_docx(self, docx_path):
        """提取信息"""
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        tables = []
        with zipfile.ZipFile(docx_path, 'r') as zf:
            xml_content = zf.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            
            for tbl in tree.findall('.//w:tbl', ns):
                table_data = []
                for row in tbl.findall('.//w:tr', ns):
                    row_data = []
                    for cell in row.findall('.//w:tc', ns):
                        cell_text = ''
                        for text in cell.findall('.//w:t', ns):
                            if text.text:
                                cell_text += text.text
                        row_data.append(cell_text.strip())
                    if row_data:
                        table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
        
        info = {'software_name': '', 'version': 'V1.0', 'completion_date': '', 'owners': []}
        
        for table in tables:
            for i in range(len(table)):
                row = table[i]
                row_text = ' '.join(row)
                
                if '软著名称' in row_text and i + 1 < len(table):
                    data = table[i + 1]
                    info['software_name'] = data[0] if data else ''
                    info['version'] = data[2] if len(data) > 2 and data[2] else 'V1.0'
                    info['completion_date'] = data[3] if len(data) > 3 else ''
                
                if '著作权人信息' in row_text:
                    for j in range(i + 1, min(i + 3, len(table))):
                        if '公司/单位' in ' '.join(table[j]):
                            for k in range(j + 1, len(table)):
                                r = table[k]
                                if '联系人' in ' '.join(r) or not r or not r[0]:
                                    break
                                if '公司' in r[0] or '营业执照' in r[0]:
                                    continue
                                id_num = r[1] if len(r) > 1 else ''
                                info['owners'].append({
                                    'name': r[0],
                                    'id_number': id_num,
                                    'is_individual': self.is_individual(id_num)
                                })
                            break
        return info
        
    def is_individual(self, id_number):
        """判断是否为个人"""
        if not id_number:
            return False
        import re
        return bool(re.match(r'^\d{17}[\dXx]$', id_number.strip()))
        
    def calc_contract_date(self, completion_str):
        """计算协议日期"""
        try:
            if '年' in completion_str:
                parts = completion_str.replace('年', ' ').replace('月', ' ').replace('日', '').split()
                y, m, d = int(parts[0]), int(parts[1]), int(parts[2])
                m = m - 3
                if m <= 0:
                    m += 12
                    y -= 1
                return f'{y}年{m}月{d}日'
        except:
            pass
        return '2025年1月1日'
        
    def generate(self):
        """生成协议"""
        if not self.input_files:
            messagebox.showwarning("提示", "请先添加申请表文件")
            return
            
        output_dir = self.output_entry.get()
        template_dir = self.template_entry.get()
        
        os.makedirs(output_dir, exist_ok=True)
        
        success_count = 0
        fail_count = 0
        errors = []
        
        self.status_var.set("正在生成...")
        self.root.update()
        
        for file_path in self.input_files:
            try:
                result = self.generate_contract(file_path, template_dir, output_dir)
                if result['success']:
                    success_count += 1
                else:
                    fail_count += 1
                    errors.append(f"{os.path.basename(file_path)}: {result['error']}")
            except Exception as e:
                fail_count += 1
                errors.append(f"{os.path.basename(file_path)}: {str(e)}")
        
        if fail_count == 0:
            msg = f"生成完成！\n\n成功: {success_count} 个\n输出目录: {output_dir}"
            messagebox.showinfo("成功", msg)
        else:
            msg = f"处理完成\n\n成功: {success_count} 个\n失败: {fail_count} 个\n\n失败原因:\n"
            msg += "\n".join(errors[:5])
            if len(errors) > 5:
                msg += f"\n...等{len(errors)}个错误"
            messagebox.showwarning("完成", msg)
            
        self.status_var.set(f"成功: {success_count}, 失败: {fail_count}")
        
    def generate_contract(self, app_path, template_dir, output_dir):
        """生成单个协议"""
        info = self.extract_info_from_docx(app_path)
        
        if not info['software_name']:
            return {'success': False, 'error': '无法提取软件名称'}
        if not info['owners']:
            return {'success': False, 'error': '无法提取著作权人'}
        
        owner_count = len(info['owners'])
        template_path = os.path.join(template_dir, f'合作协议-{owner_count}方.docx')
        
        if not os.path.exists(template_path):
            return {'success': False, 'error': f'模板不存在: {owner_count}方.docx'}
        
        contract_date = self.calc_contract_date(info['completion_date'])
        software_full = f"{info['software_name']}{info['version']}"
        copy_count = owner_count + 1
        
        doc = Document(template_path)
        
        replacements = [
            ('智能心理压力情绪疏导与调解平台V1.0', software_full),
            ('2025-3-10', contract_date),
        ]
        
        sample_names = ['王红梅', '汪庆军', '吴六韬', '赖羽羽', '刘茜茜',
                        '丁姿慧', '刘澈', '曹冬梅', '谢奇秀', '钟海玲']
        for i, owner in enumerate(info['owners']):
            if i < len(sample_names):
                replacements.append((sample_names[i], owner['name']))
        
        for nc in ['三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二', '十三', '十四']:
            replacements.append((f'{nc}份', f'{copy_count}份'))
        
        for old_text, new_text in replacements:
            for para in doc.paragraphs:
                self.replace_in_runs(para, old_text, new_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self.replace_in_runs(para, old_text, new_text)
        
        self.add_signature_names(doc, info['owners'])
        
        safe_name = info['software_name'].replace('/', '-').replace('\\', '-').replace(':', '-')
        output_path = os.path.join(output_dir, f'合作协议-{safe_name}.docx')
        doc.save(output_path)
        
        return {'success': True, 'output_file': output_path}
        
    def replace_in_runs(self, paragraph, old_text, new_text):
        if old_text not in paragraph.text:
            return
        runs = paragraph.runs
        if not runs:
            return
        new_full_text = paragraph.text.replace(old_text, new_text)
        for run in runs:
            run.text = ''
        runs[0].text = new_full_text
        
    def add_signature_names(self, doc, owners):
        body = doc.element.body
        paras = body.findall(qn('w:p'))
        
        signature_start = -1
        found_yishi = False
        for i, para in enumerate(paras):
            text = ''.join(t.text for t in para.findall('.//' + qn('w:t')) if t.text)
            if '一式' in text:
                found_yishi = True
            if found_yishi and text.strip() == '甲方：':
                signature_start = i
                break
        
        if signature_start < 0:
            return
        
        party_names = ['甲方：', '乙方：', '丙方：', '丁方：', '戊方：', '己方：', 
                       '庚方：', '辛方：', '壬方：', '癸方：', '第十一方：', '第十二方：', '第十三方：']
        
        for i, para in enumerate(paras[signature_start:], start=signature_start):
            text = ''.join(t.text for t in para.findall('.//' + qn('w:t')) if t.text)
            text = text.strip()
            
            if text in party_names:
                party_idx = party_names.index(text)
                if party_idx < len(owners):
                    owner = owners[party_idx]
                    if not owner.get('is_individual'):
                        runs = para.findall('.//' + qn('w:r'))
                        if runs:
                            new_run = OxmlElement('w:r')
                            rPr = runs[0].find(qn('w:rPr'))
                            if rPr is not None:
                                new_run.append(deepcopy(rPr))
                            new_text = OxmlElement('w:t')
                            new_text.text = owner['name']
                            new_run.append(new_text)
                            para.append(new_run)


def main():
    root = Tk()
    app = ContractGenerator(root)
    root.mainloop()


if __name__ == '__main__':
    main()
